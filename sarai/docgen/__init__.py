"""MinutesJSON -> .docx.

Two rules run through everything below.

*Every* piece of text goes through `styled_run`. python-docx sets `w:ascii` and
`w:hAnsi` when you assign `run.font.name`, and stops there -- Thai is a complex
script, so without `w:cs` Word picks its own substitute font, and without
`w:szCs` it renders at the wrong point size. Missing either is invisible on a
developer machine with Sarabun installed and obvious to everyone else.

The section order is the one Thai minutes are read in: attendees, executive
summary, agenda, discussion, decisions, action items, open questions. Action
items are the section people actually act on, so they get a table.
"""

from __future__ import annotations

import logging
from datetime import UTC, datetime
from pathlib import Path

from docx import Document as new_document
from docx.document import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from sarai.docgen.fonts import FONT_NAME, embed_fonts
from sarai.models import MinutesJSON

log = logging.getLogger("sarai.docgen")

BODY_PT = 12
H1_PT = 18
H2_PT = 14
SMALL_PT = 10


def styled_run(
    paragraph: Paragraph,
    text: str,
    *,
    size: int = BODY_PT,
    bold: bool = False,
    italic: bool = False,
    color: tuple[int, int, int] | None = None,
) -> Run:
    """The only way text enters the document. See the module docstring."""
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor(*color)

    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), FONT_NAME)
    rfonts.set(qn("w:hAnsi"), FONT_NAME)
    rfonts.set(qn("w:cs"), FONT_NAME)  # complex script -- required for Thai
    # w:sz is set by python-docx above; w:szCs is the complex-script twin and
    # is what actually sizes Thai glyphs. Half-points, hence the doubling.
    sz_cs = rpr.makeelement(qn("w:szCs"), {qn("w:val"): str(size * 2)})
    rpr.append(sz_cs)
    return run


def _heading(doc: Document, text: str, *, size: int = H2_PT, space_before: int = 14) -> Paragraph:
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(4)
    styled_run(para, text, size=size, bold=True)
    return para


def _body(doc: Document, text: str, *, size: int = BODY_PT, italic: bool = False) -> Paragraph:
    para = doc.add_paragraph()
    # Thai tone marks collide with the line above at the default spacing.
    para.paragraph_format.line_spacing = 1.5
    para.paragraph_format.space_after = Pt(4)
    styled_run(para, text, size=size, italic=italic)
    return para


def _bullet(doc: Document, text: str, *, prefix: str = "•") -> Paragraph:
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Pt(18)
    para.paragraph_format.line_spacing = 1.5
    para.paragraph_format.space_after = Pt(2)
    styled_run(para, f"{prefix} {text}")
    return para


def _bilingual(thai_label: str, english_label: str) -> str:
    return f"{thai_label} / {english_label}"


def _attendee_table(doc: Document, minutes: MinutesJSON) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for attendee in minutes.attendees:
        row = table.add_row().cells
        styled_run(row[0].paragraphs[0], attendee.name)
        styled_run(row[1].paragraphs[0], attendee.role or "—")


def _action_table(doc: Document, minutes: MinutesJSON) -> None:
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    header = table.rows[0].cells
    for cell, label in zip(
        header,
        (
            _bilingual("สิ่งที่ต้องทำ", "Task"),
            _bilingual("ผู้รับผิดชอบ", "Owner"),
            _bilingual("กำหนดส่ง", "Due"),
        ),
        strict=True,
    ):
        styled_run(cell.paragraphs[0], label, bold=True, size=SMALL_PT + 1)

    for item in minutes.action_items:
        cells = table.add_row().cells
        styled_run(cells[0].paragraphs[0], item.task)
        styled_run(cells[1].paragraphs[0], item.owner or "—")
        styled_run(cells[2].paragraphs[0], item.due or "—")
        # The quote is the evidence the item is real; it belongs next to it.
        quote = cells[0].add_paragraph()
        styled_run(
            quote, f"“{item.source_quote}”", size=SMALL_PT, italic=True, color=(0x6F, 0x66, 0x59)
        )


def render(minutes: MinutesJSON, dest: Path, *, model: str) -> Path:
    """Write the minutes to `dest`. Returns the path written."""
    dest.parent.mkdir(parents=True, exist_ok=True)
    doc = new_document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    styled_run(title, minutes.title, size=H1_PT, bold=True)

    generated = datetime.now(UTC).strftime("%Y-%m-%d %H:%M UTC")
    meta = doc.add_paragraph()
    styled_run(
        meta,
        f"{_bilingual('วันที่ประชุม', 'Date')}: {minutes.meeting_date or '—'}",
        size=SMALL_PT,
    )

    if minutes.attendees:
        _heading(doc, _bilingual("ผู้เข้าร่วมประชุม", "Attendees"))
        _attendee_table(doc, minutes)

    if minutes.summary:
        _heading(doc, _bilingual("สรุปผู้บริหาร", "Executive summary"))
        _body(doc, minutes.summary)

    if minutes.agenda:
        _heading(doc, _bilingual("วาระการประชุม", "Agenda"))
        for index, item in enumerate(minutes.agenda, start=1):
            _bullet(doc, item, prefix=f"{index}.")

    if minutes.discussion:
        _heading(doc, _bilingual("รายละเอียดการประชุม", "Discussion"))
        for topic in minutes.discussion:
            _heading(doc, topic.topic, size=BODY_PT + 1, space_before=8)
            if topic.speakers:
                _body(doc, ", ".join(topic.speakers), size=SMALL_PT, italic=True)
            for point in topic.points:
                _bullet(doc, point)

    if minutes.decisions:
        _heading(doc, _bilingual("มติที่ประชุม", "Decisions"))
        for index, decision in enumerate(minutes.decisions, start=1):
            _bullet(doc, decision.decision, prefix=f"{index}.")
            if decision.rationale:
                para = doc.add_paragraph()
                para.paragraph_format.left_indent = Pt(36)
                styled_run(para, decision.rationale, size=SMALL_PT, italic=True)

    if minutes.action_items:
        # The section people actually read; a long document should not make
        # them hunt for it across a page break.
        doc.add_page_break()  # type: ignore[no-untyped-call]
        _heading(doc, _bilingual("สิ่งที่ต้องดำเนินการ", "Action items"), size=H2_PT + 1)
        _action_table(doc, minutes)

    if minutes.open_questions:
        _heading(doc, _bilingual("ประเด็นค้าง", "Open questions"))
        for question in minutes.open_questions:
            _bullet(doc, question)

    if minutes.next_meeting:
        _heading(doc, _bilingual("การประชุมครั้งต่อไป", "Next meeting"))
        _body(doc, minutes.next_meeting)

    footer = doc.add_paragraph()
    footer.paragraph_format.space_before = Pt(24)
    styled_run(
        footer,
        f"Generated by Sarai AI · {model} · {generated}",
        size=SMALL_PT,
        color=(0x9C, 0x91, 0x84),
    )

    embed_fonts(doc)
    doc.save(str(dest))
    log.info("rendered minutes to %s", dest)
    return dest
